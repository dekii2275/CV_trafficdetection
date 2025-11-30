"""
Script ƒë·ªÉ build vector database t·ª´ c√°c file lu·∫≠t giao th√¥ng.
ƒê√£ s·ª≠a l·ªói Import LangChain v√† l·ªói tham s·ªë d√≤ng l·ªánh.

"""

import sys
import os
import re
import shutil
from typing import List
from pathlib import Path
import docx  

try:
    
    from langchain_core.documents import Document
except ImportError:
    try:
        from langchain.schema import Document
    except ImportError:
        from langchain.docstore.document import Document


FILE_PATH = Path(__file__).resolve()

PROJECT_ROOT = FILE_PATH.parents[3] 
BACKEND_ROOT = FILE_PATH.parents[2] 


sys.path.append(str(BACKEND_ROOT))

ABS_DOCS_DIR = PROJECT_ROOT / "data" / "law_documents"
ABS_DB_DIR = PROJECT_ROOT / "data" / "chroma_db"

# Import service c·ªßa b·∫°n
try:
    from app.services.rag_services.vector_store import VectorStoreService
except ImportError as e:
    print(f"L·ªói Import Service: {e}")
    print(f"ƒê·∫£m b·∫£o b·∫°n ƒëang ƒë·ª©ng ·ªü th∆∞ m·ª•c 'backend' v√† file vector_store.py t·ªìn t·∫°i.")
    sys.exit(1)


class TrafficLawProcessor:
    """
    X·ª≠ l√Ω vƒÉn b·∫£n lu·∫≠t: T√°ch ƒêi·ªÅu -> Kho·∫£n -> ƒêi·ªÉm ƒë·ªÉ tr√°nh m·∫•t th√¥ng tin
    """
    def __init__(self):
        # Regex t√¨m "ƒêi·ªÅu X."
        self.article_pattern = r"(^|\n)(ƒêi·ªÅu \d+\..*?)(?=\nƒêi·ªÅu \d+\.|$)"
        # Regex t√¨m "1. ", "2. " (Kho·∫£n)
        self.clause_pattern = r"(^|\n)(\d+)\.\s+(.*?)(?=(\n\d+\.\s+)|$)"
        # Regex t√¨m "a) ", "b) ", "ƒë) " (ƒêi·ªÉm)
        self.point_pattern = r"(^|\n)([a-zƒë])\)\s+(.*?)(?=(\n[a-zƒë]\))|$)" 
    
    def read_docx(self, file_path: str) -> str:
        """ƒê·ªçc file .docx v√† chuy·ªÉn th√†nh string"""
        try:
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                txt = para.text.strip()
                if txt:
                    full_text.append(txt)
            return "\n".join(full_text)
        except Exception as e:
            print(f"L·ªói ƒë·ªçc file {file_path}: {e}")
            return ""

    def identify_vehicle_type(self, text: str) -> str:
        """Nh·∫≠n di·ªán lo·∫°i xe t·ª´ ti√™u ƒë·ªÅ ƒêi·ªÅu lu·∫≠t"""
        text_lower = text.lower()
        if "xe √¥ t√¥" in text_lower: return "oto"
        if "xe m√¥ t√¥" in text_lower or "xe g·∫Øn m√°y" in text_lower: return "xemay"
        if "xe ƒë·∫°p" in text_lower or "xe th√¥ s∆°" in text_lower: return "xedap"
        if "ng∆∞·ªùi ƒëi b·ªô" in text_lower: return "nguoidibo"
        return "chung"

    def process_document(self, file_path: str) -> List[Document]:
        text = self.read_docx(file_path)
        if not text: return []
        
        chunks = []
        source_name = Path(file_path).name
        
        # B1: T√°ch c√°c ƒêi·ªÅu
        articles = re.finditer(self.article_pattern, text, re.DOTALL)
        
        for art_match in articles:
            article_full_text = art_match.group(2)
            # L·∫•y d√≤ng ƒë·∫ßu ti√™n l√†m ti√™u ƒë·ªÅ (VD: ƒêi·ªÅu 5. X·ª≠ ph·∫°t...)
            article_header = article_full_text.strip().split('\n')[0]
            vehicle_type = self.identify_vehicle_type(article_header)
            
            # B2: T√°ch c√°c Kho·∫£n (Clauses) trong ƒêi·ªÅu
            clauses = re.finditer(self.clause_pattern, article_full_text, re.DOTALL)
            has_clauses = False
            
            for clause_match in clauses:
                has_clauses = True
                clause_num = clause_match.group(2)
                clause_body = clause_match.group(3).strip()
                
                # B3: T√°ch c√°c ƒêi·ªÉm (Points: a, b, c...) trong Kho·∫£n
                points = list(re.finditer(self.point_pattern, clause_body, re.DOTALL))
                
                if points:
                    # L·∫•y ph·∫ßn d·∫´n nh·∫≠p (VD: "Ph·∫°t ti·ªÅn t·ª´ 200k... h√†nh vi sau:")
                    intro_text = clause_body[:points[0].start()].strip()
                    
                    for p_match in points:
                        p_label = p_match.group(2) # a, b, c
                        p_content = p_match.group(3).strip()
                        
                        # T·∫°o n·ªôi dung Chunk chi ti·∫øt
                        full_content = (
                            f"ƒêI·ªÄU LU·∫¨T: {article_header}\n"
                            f"M·ª®C PH·∫†T (Kho·∫£n {clause_num}): {intro_text}\n"
                            f"H√ÄNH VI VI PH·∫†M (ƒêi·ªÉm {p_label}): {p_content}"
                        )
                        
                        chunks.append(Document(
                            page_content=full_content,
                            metadata={
                                "source": source_name,
                                "article": article_header.split('.')[0], # VD: ƒêi·ªÅu 5
                                "vehicle": vehicle_type,
                                "level": "point" # C·∫•p ƒë·ªô chi ti·∫øt nh·∫•t
                            }
                        ))
                else:
                    # N·∫øu kh√¥ng c√≥ ƒëi·ªÉm a,b,c -> L·∫•y nguy√™n Kho·∫£n
                    full_content = (
                        f"ƒêI·ªÄU LU·∫¨T: {article_header}\n"
                        f"N·ªòI DUNG (Kho·∫£n {clause_num}): {clause_body}"
                    )
                    chunks.append(Document(
                        page_content=full_content,
                        metadata={
                            "source": source_name,
                            "article": article_header.split('.')[0],
                            "vehicle": vehicle_type,
                            "level": "clause"
                        }
                    ))

            # N·∫øu ƒêi·ªÅu qu√° ng·∫Øn kh√¥ng c√≥ kho·∫£n (ch·ªâ c√≥ text)
            if not has_clauses:
                chunks.append(Document(
                    page_content=article_full_text, 
                    metadata={"source": source_name, "vehicle": vehicle_type, "level": "article"}
                ))
                
        return chunks

def build_vector_database(documents_dir: str = str(ABS_DOCS_DIR), reset: bool = False):
    print("\n" + "="*60)
    print("RAG BUILDER: SMART CHUNKING (Lu·∫≠t Giao Th√¥ng)")
    print("="*60)
    print(f"ƒê·ªçc t√†i li·ªáu t·ª´: {documents_dir}")
    print(f"L∆∞u Database t·∫°i: {ABS_DB_DIR}")
    
    # X·ª≠ l√Ω tham s·ªë Reset
    if reset:
        if ABS_DB_DIR.exists():
            print(f"ƒêang x√≥a database c≈© ƒë·ªÉ l√†m s·∫°ch d·ªØ li·ªáu...")
            shutil.rmtree(ABS_DB_DIR)
        else:
            print("Kh√¥ng t√¨m th·∫•y database c≈©, s·∫Ω t·∫°o m·ªõi ho√†n to√†n.")
    
    # Init Vector Store
    print("ƒêang kh·ªüi t·∫°o Vector Store...")
    vector_store = VectorStoreService(
        collection_name="traffic_laws",
        persist_directory=str(ABS_DB_DIR)
    )

    # X·ª≠ l√Ω file
    processor = TrafficLawProcessor()
    all_documents = []
    
    if not os.path.exists(documents_dir):
        print(f"L·ªñI: Th∆∞ m·ª•c t√†i li·ªáu kh√¥ng t·ªìn t·∫°i: {documents_dir}")
        print(f"Vui l√≤ng t·∫°o th∆∞ m·ª•c n√†y v√† copy file .docx v√†o ƒë√≥.")
        return

    files = [f for f in os.listdir(documents_dir) if f.endswith(".docx") or f.endswith(".doc")]
    if not files:
        print(f"C·∫¢NH B√ÅO: Th∆∞ m·ª•c {documents_dir} tr·ªëng! H√£y copy file lu·∫≠t v√†o.")
        return

    for filename in files:
        file_path = os.path.join(documents_dir, filename)
        print(f"\nƒêang x·ª≠ l√Ω file: {filename}...")
        
        chunks = processor.process_document(file_path)
        all_documents.extend(chunks)
        print(f"   -> T·∫°o ƒë∆∞·ª£c {len(chunks)} chunks d·ªØ li·ªáu.")

    if not all_documents:
        print(" Kh√¥ng t·∫°o ƒë∆∞·ª£c d·ªØ li·ªáu n√†o. Ki·ªÉm tra l·∫°i n·ªôi dung file input.")
        return

    # L∆∞u v√†o ChromaDB
    print(f"\n ƒêang l∆∞u {len(all_documents)} chunks v√†o Database (Qu√° tr√¨nh n√†y c√≥ th·ªÉ m·∫•t v√†i ph√∫t)...")
    
    texts = [doc.page_content for doc in all_documents]
    metadatas = [doc.metadata for doc in all_documents]
    
    vector_store.add_documents(documents=texts, metadatas=metadatas)
    print("\nX√ÇY D·ª∞NG DATABASE TH√ÄNH C√îNG!")


# 5. H√ÄM TEST TRUY V·∫§N

def test_search(query: str):
    print("\n" + "="*60)
    print(f"TEST TRUY V·∫§N TH·ª¨: \"{query}\"")
    print("="*60)
    
    if not ABS_DB_DIR.exists():
        print("Database ch∆∞a ƒë∆∞·ª£c x√¢y d·ª±ng. H√£y ch·∫°y l·ªánh build tr∆∞·ªõc.")
        return

    vector_store = VectorStoreService(
        collection_name="traffic_laws",
        persist_directory=str(ABS_DB_DIR)
    )
    
    results = vector_store.search(query, top_k=3)
    
    print(f"üîç T√¨m th·∫•y {len(results)} k·∫øt qu·∫£ li√™n quan nh·∫•t:\n")
    for i, res in enumerate(results, 1):
        # L·∫•y th√¥ng tin an to√†n
        score = res.get('similarity_score', 0)
        meta = res.get('metadata', {})
        content = res.get('document', '')
        
        # L√†m g·ªçn n·ªôi dung ƒë·ªÉ hi·ªÉn th·ªã
        preview = content[:250].replace('\n', ' ') + "..."
        
        print(f"--- Top {i} (ƒê·ªô kh·ªõp: {score:.2f}) ---")
        print(f"Ngu·ªìn: {meta.get('source', 'N/A')} | {meta.get('article', 'N/A')}")
        print(f"N·ªôi dung: {preview}\n")

if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Tool build d·ªØ li·ªáu cho Chatbot Giao th√¥ng")
    
    parser.add_argument("--reset", action="store_true", help="X√≥a s·∫°ch DB c≈© v√† build l·∫°i t·ª´ ƒë·∫ßu")
    
    parser.add_argument("--test-query", type=str, default="kh√¥ng ƒë·ªôi m≈© b·∫£o hi·ªÉm ph·∫°t bao nhi√™u", help="C√¢u h·ªèi ƒë·ªÉ test th·ª≠ sau khi build")
    
    parser.add_argument("--skip-build", action="store_true", help="Ch·ªâ ch·∫°y test, kh√¥ng build l·∫°i DB")
    
    args = parser.parse_args()
    
    if not args.skip_build:
        build_vector_database(reset=args.reset)
    
    test_search(args.test_query)